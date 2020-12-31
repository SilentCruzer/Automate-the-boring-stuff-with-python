import docx

doc = docx.Document('invitations.docx')
name_file = open('name_lists.txt','r')
names = name_file.readlines()
names = [name .strip() for name in names]

#Custom styles should be saved in an docx file before running the program
for name in names:
	doc.add_paragraph('It would be a pleasure to have company of').style = 'introline'
	doc.add_paragraph(name).style = 'name'
	doc.add_paragraph('at 11010 Memory lane in the evening of').style = 'introline'
	doc.add_paragraph('April 1st').style = 'name'
	end = doc.add_paragraph("at 7 o'clock").style = 'introline'
	end.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)



